from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
INPUT = ROOT / "调研报告-初稿.docx"
OUTPUT = ROOT / "数字媒体技术课程调研报告-朱清扬-李奕霖.docx"

doc = Document(INPUT)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_widths(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = int(sum(widths_cm) / 2.54 * 1440)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_cm)):
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=12, bold=None, color=None):
    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, east_asia="宋体", size=9)


section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.6)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)
footer = section.footer
footer.paragraphs[0].clear()
add_page_number(footer.paragraphs[0])

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = Pt(22)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)

for paragraph in list(doc.paragraphs):
    if paragraph.text.strip().startswith("{width="):
        for run in paragraph.runs:
            if "{width=" in run.text:
                run.text = ""

for table in doc.tables:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        for height in list(tr_pr.findall(qn("w:trHeight"))):
            tr_pr.remove(height)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if "{width=" in run.text:
                        run.text = run.text.replace("{width=1.6in}", "")

major_headings = {"调研报告摘要", "调研报告正文", "参考文献", "调研报告成绩评定表"}
section_headings = {
    "引言与问题界定",
    "产业背景与产品演进",
    "学术研究与方法证据",
    "开源生态与工业佐证",
    "多智能体创作流程重构",
    "关键挑战与未来趋势",
    "结论",
}
captions = {
    "图1. Google Flow 将生成模型组织为面向镜头与场景的创作工具",
    "图2. Adobe Firefly Video 面向生成、编辑与后期衔接的一体化界面",
    "图3. MM-StoryAgent 将文本、图像、语音、音效和音乐智能体组织为视频生产管线",
    "表1. 典型 AI 视频产品的流程化能力比较",
}

in_references = False
for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text == "参考文献":
        in_references = True
    elif text == "调研报告成绩评定表":
        in_references = False
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(22)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.widow_control = True

    if index == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_before = Pt(60)
        fmt.space_after = Pt(22)
        for run in paragraph.runs:
            set_run_font(run, east_asia="黑体", latin="Arial", size=20, bold=True)
    elif index == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.line_spacing = Pt(28)
        fmt.space_after = Pt(28)
        for run in paragraph.runs:
            set_run_font(run, east_asia="黑体", latin="Arial", size=18, bold=True)
    elif text in major_headings:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.page_break_before = True
        fmt.keep_with_next = True
        fmt.space_after = Pt(12)
        for run in paragraph.runs:
            set_run_font(run, east_asia="黑体", latin="Arial", size=16, bold=True)
    elif text in section_headings:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.keep_with_next = True
        fmt.space_before = Pt(8)
        fmt.space_after = Pt(4)
        for run in paragraph.runs:
            set_run_font(run, east_asia="黑体", latin="Arial", size=14, bold=True)
    elif text in captions:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.keep_with_next = text.startswith("表")
        fmt.space_before = Pt(3)
        fmt.space_after = Pt(5)
        fmt.line_spacing = Pt(18)
        for run in paragraph.runs:
            set_run_font(run, east_asia="宋体", size=10.5, bold=True)
    elif text.startswith("关键词：") or text == "调研团队成员电子签名：":
        fmt.space_before = Pt(4)
        fmt.space_after = Pt(4)
        for run in paragraph.runs:
            set_run_font(run, size=12, bold=run.bold)
    elif text.startswith("指导教师签章："):
        fmt.space_before = Pt(16)
        for run in paragraph.runs:
            set_run_font(run, size=12)
    elif in_references:
        fmt.left_indent = Cm(0.75)
        fmt.first_line_indent = Cm(-0.75)
        fmt.line_spacing = Pt(18)
        fmt.space_after = Pt(3)
        for run in paragraph.runs:
            set_run_font(run, east_asia="宋体", latin="Times New Roman", size=10.5)
    elif text:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Pt(24)
        for run in paragraph.runs:
            set_run_font(run, size=12, bold=run.bold)

for i, shape in enumerate(doc.inline_shapes):
    if i < 2:
        shape.width = Inches(1.55)
        shape.height = Inches(0.44)
    elif i == 2:
        shape.width = Inches(5.55)
        shape.height = Inches(3.12)
    elif i == 3:
        shape.width = Inches(5.55)
        shape.height = Inches(3.11)
    else:
        shape.width = Inches(5.7)
        shape.height = Inches(3.68)
    parent = shape._inline.getparent()
    while parent is not None and parent.tag != qn("w:p"):
        parent = parent.getparent()
    if parent is not None:
        p_pr = parent.get_or_add_pPr()
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        for paragraph in doc.paragraphs:
            if paragraph._p is parent:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.keep_with_next = True
                break

cover_table = doc.tables[0]
set_table_widths(cover_table, [3.2, 10.8])
for r_idx, row in enumerate(cover_table.rows):
    for c_idx, cell in enumerate(row.cells):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=150, bottom=150)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(20)
            for run in p.runs:
                set_run_font(run, east_asia="宋体", size=12, bold=(r_idx == 0 or c_idx == 0))
        if r_idx == 0:
            set_cell_shading(cell, "D9EAF7")

signature_table = doc.tables[1]
set_table_widths(signature_table, [7.0, 7.0])
signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in signature_table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=40, bottom=40)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = Pt(18)
            for run in p.runs:
                set_run_font(run, size=10.5)

comparison = doc.tables[2]
set_table_widths(comparison, [3.0, 5.0, 6.0])
set_repeat_table_header(comparison.rows[0])
for r_idx, row in enumerate(comparison.rows):
    for c_idx, cell in enumerate(row.cells):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=100, bottom=100)
        if r_idx == 0:
            set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 or r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(18)
            for run in p.runs:
                set_run_font(run, size=10.5, bold=(r_idx == 0))

score = doc.tables[3]
set_table_widths(score, [2.0, 8.5, 2.2, 1.3])
set_repeat_table_header(score.rows[0])
for r_idx, row in enumerate(score.rows):
    for c_idx, cell in enumerate(row.cells):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=130, bottom=130)
        if r_idx == 0:
            set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 or r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(19)
            for run in p.runs:
                set_run_font(run, size=10.5, bold=(r_idx == 0))

doc.core_properties.title = "从文本到短片：多智能体协同在 AI 视频生成与数字媒体创作流程中的应用调研"
doc.core_properties.subject = "2026年《数字媒体技术》课程调研报告"
doc.core_properties.author = "朱清扬、李奕霖"
doc.core_properties.keywords = "AI视频生成；多智能体系统；数字媒体创作；工作流；人机协同"

doc.save(OUTPUT)
print(OUTPUT)
