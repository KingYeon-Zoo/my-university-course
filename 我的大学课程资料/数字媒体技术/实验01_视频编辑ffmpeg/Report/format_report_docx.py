from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "数媒2026-实验报告01-朱清扬-2023212290.docx"


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=None, bold=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_margins(table, top=80, start=100, bottom=80, end=100):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches):
    widths_twips = [int(width * 1440) for width in widths_inches]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblBorders"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "7F8C8D")
        borders.append(border)
    tbl_pr.append(borders)

    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for column_index, cell in enumerate(row.cells):
            width = widths_twips[column_index]
            cell.width = Inches(widths_inches[column_index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_cell_margins(table)


def style_table(table, widths, font_size):
    set_table_geometry(table, widths)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    header_row = table.rows[0]
    tr_pr = header_row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)

    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                if row_index == 0 or column_index in (0, 1):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(row_index == 0))
            if row_index == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D9EAF7")
                tc_pr.append(shading)


doc = Document(DOCX)

heading_texts = {
    "实验环境与规范",
    "视频编码及处理操作完成情况",
    "操作过程及命令",
    "最终视频交付检查",
    "其他说明",
    "实验报告成绩评定表",
}

for index, paragraph in enumerate(doc.paragraphs):
    if index == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(12)
        for run in paragraph.runs:
            set_run_font(run, east_asia="黑体", size=16, bold=True)
    elif index == 1:
        paragraph.paragraph_format.space_after = Pt(10)
        for run in paragraph.runs:
            set_run_font(run, east_asia="宋体", size=11, bold=True)

    text = paragraph.text.strip()
    if text in heading_texts:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
        if text == "实验报告成绩评定表":
            paragraph.paragraph_format.page_break_before = True
        for run in paragraph.runs:
            set_run_font(run, east_asia="黑体", size=13, bold=True)

    if paragraph._p.xpath(".//w:drawing"):
        p_pr = paragraph._p.get_or_add_pPr()
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            p_pr.remove(p_style)
        if "指导教师签章" not in text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(6)

style_table(doc.tables[0], [1.30, 0.70, 4.20], 8.5)
style_table(doc.tables[1], [0.65, 3.50, 1.00, 0.85], 9)

target_widths = [6.10, 6.10, 6.10, 0.86]
for shape, width in zip(doc.inline_shapes, target_widths):
    ratio = shape.height / shape.width
    shape.width = Inches(width)
    shape.height = int(shape.width * ratio)

doc.save(DOCX)
