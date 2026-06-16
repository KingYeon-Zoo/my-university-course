from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 1. 生成模板 ---------- #
def create_template(path):
    doc = Document()
    # 版心与页面设置
    s = doc.sections[0]
    s.left_margin, s.right_margin = Cm(2.8), Cm(2.8)
    s.top_margin,  s.bottom_margin = Cm(3),  Cm(3)
    s.header_distance = s.footer_distance = Cm(1.5)

    # 正文样式 - 符合规范：宋体小四号、22磅行距、首行缩进2字符、禁用彩色文字
    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # 数字字母使用Times New Roman
    normal.font.size = Pt(12)  # 小四号
    normal.font.color.rgb = RGBColor(0, 0, 0)  # 明确设置黑色，禁用彩色文字
    normal.font.italic = False  # 禁用斜体，使用正体文字
    pf = normal.paragraph_format
    pf.first_line_indent = Pt(24)  # 首行缩进2字符
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    pf.line_spacing_rule, pf.line_spacing = WD_LINE_SPACING.EXACTLY, Pt(22)  # 22磅固定行距

    # 一级标题 - 符合规范：黑体三号加粗居中、段前段后各1行间距、禁用彩色文字
    h1 = doc.styles['Heading 1']
    h1.font.name = '黑体'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h1._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # 数字字母使用Times New Roman
    h1.font.bold, h1.font.size = True, Pt(16)  # 黑体三号加粗
    h1.font.color.rgb = RGBColor(0, 0, 0)  # 明确设置黑色，禁用彩色文字
    h1.font.italic = False  # 禁用斜体，使用正体文字
    h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    h1.paragraph_format.space_before = Pt(22)  # 段前1行间距
    h1.paragraph_format.space_after = Pt(22)   # 段后1行间距

    # 二级标题 - 符合规范：黑体小四号左对齐、段前段后各0.5行间距、禁用彩色文字
    h2 = doc.styles['Heading 2']
    h2.font.name = '黑体'
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h2._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # 数字字母使用Times New Roman
    h2.font.bold, h2.font.size = True, Pt(12)  # 黑体小四号加粗
    h2.font.color.rgb = RGBColor(0, 0, 0)  # 明确设置黑色，禁用彩色文字
    h2.font.italic = False  # 禁用斜体，使用正体文字
    h2.paragraph_format.space_before = Pt(11)  # 段前0.5行间距
    h2.paragraph_format.space_after = Pt(11)   # 段后0.5行间距

    # 三级标题 - 符合规范：宋体小四号加粗左对齐、段前段后各0.5行间距、禁用彩色文字
    h3 = doc.styles['Heading 3']
    h3.font.name = '宋体'
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h3._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # 数字字母使用Times New Roman
    h3.font.bold, h3.font.size = True, Pt(12)  # 宋体小四号加粗
    h3.font.color.rgb = RGBColor(0, 0, 0)  # 明确设置黑色，禁用彩色文字
    h3.font.italic = False  # 禁用斜体，使用正体文字
    h3.paragraph_format.space_before = Pt(11)  # 段前0.5行间距
    h3.paragraph_format.space_after = Pt(11)   # 段后0.5行间距

    # 页眉 - 符合规范：宋体五号居中、禁用彩色文字
    p = s.header.paragraphs[0]
    run = p.add_run('合肥工业大学本科毕业设计（论文）')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # 数字字母使用Times New Roman
    run.font.size = Pt(10.5)  # 五号字体
    run.font.color.rgb = RGBColor(0, 0, 0)  # 明确设置黑色，禁用彩色文字
    run.font.italic = False  # 禁用斜体，使用正体文字
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐

    # 页码
    p = s.footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    for typ in ('begin', 'separate', 'end'):
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), typ)
        run._r.append(fld) if typ != 'separate' else None
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    run._r.insert(1, instr)
    doc.save(path)

# ---------- 2. Markdown 转 Word ---------- #
def markdown_to_docx(md_text, template_path, output_path):
    doc = Document(template_path)
    
    for raw in md_text.splitlines():
        line = raw.strip()
        
        # 处理标题
        if line.startswith('# '):
            doc.add_paragraph(line[2:].strip(), style='Heading 1')
        elif line.startswith('## '):
            doc.add_paragraph(line[3:].strip(), style='Heading 2')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:].strip(), style='Heading 3')
        elif line.startswith('#### '):
            # 四级标题用三级标题样式
            doc.add_paragraph(line[5:].strip(), style='Heading 3')
        # 处理空行
        elif line == '':
            doc.add_paragraph('')
        # 处理表格分隔符（跳过markdown表格分隔符行）
        elif line.startswith('|') and (':---' in line or '---' in line):
            continue
        # 处理表格行
        elif line.startswith('| ') and line.endswith(' |'):
            # 简单处理：将表格行转为普通段落
            table_content = line[1:-1].strip()  # 去除首尾的|
            cells = [cell.strip() for cell in table_content.split('|')]
            
            # 如果看起来是表头（包含中文冒号或"场景"等关键词），加粗
            if any(keyword in table_content for keyword in ['场景', '评估', '指标', '维度', '描述']):
                p = doc.add_paragraph()
                for i, cell in enumerate(cells):
                    if i > 0:
                        p.add_run('  |  ')
                    run = p.add_run(cell)
                    run.bold = True
            else:
                p = doc.add_paragraph()
                for i, cell in enumerate(cells):
                    if i > 0:
                        p.add_run('  |  ')
                    p.add_run(cell)
        # 处理图片引用（转为文本说明）
        elif line.startswith('![') and '](' in line and line.endswith(')'):
            # 提取图片描述和路径
            desc_end = line.find('](')
            if desc_end > 2:
                desc = line[2:desc_end]
                path = line[desc_end+2:-1]
                p = doc.add_paragraph()
                p.add_run('【图片】').bold = True
                p.add_run(f' {desc}（文件：{path}）')
        # 处理引用块
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Normal'
            # 增加左缩进表示引用
            p.paragraph_format.left_indent = Pt(48)  # 额外缩进
        # 处理普通段落（包含粗体等格式）
        else:
            p = doc.add_paragraph(style='Normal')
            
            # 处理粗体文本
            import re
            parts = re.split(r'(\*\*.*?\*\*)', line)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**') and len(part) > 4:
                    # 粗体文本
                    bold_text = part[2:-2]
                    run = p.add_run(bold_text)
                    run.bold = True
                elif part:
                    # 普通文本
                    p.add_run(part)
    
    doc.save(output_path)

# ---------- 3. 主程序：转换课程论文 ---------- #
if __name__ == "__main__":
    # 创建模板
    print("正在创建Word模板...")
    create_template('hfut_template.docx')
    
    # 读取课程论文markdown文件
    print("正在读取课程论文...")
    try:
        with open('课程论文.md', 'r', encoding='utf-8') as f:
            paper_content = f.read()
        
        # 转换为Word文档
        print("正在转换为Word文档...")
        markdown_to_docx(paper_content, 'hfut_template.docx', '课程论文.docx')
        
        print("转换完成！输出文件：课程论文.docx")
        print("注意：图片文件需要手动插入到Word文档中相应位置")
        
    except FileNotFoundError:
        print("错误：找不到文件 '课程论文.md'")
    except Exception as e:
        print(f"转换过程中发生错误：{e}")

    # 保留原有示例代码（注释掉）
    """
    # 原有示例代码
    sample_md = '''
    # 第一章 绪论
    示例正文，验证行距与缩进。
    ## 1.1 研究背景
    ### 1.1.1 国内研究现状
    '''
    markdown_to_docx(sample_md, 'hfut_template.docx', 'hfut_demo.docx')
    """
