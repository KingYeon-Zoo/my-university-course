from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path

from docx import Document
from PIL import Image


ROOT = Path("/Users/zoo/Desktop/计算机实训")


def docx_text(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        tables.append([[cell.text.strip() for cell in row.cells] for row in table.rows])
    return {"paragraphs": paragraphs, "tables": tables}


def image_info(path: Path) -> dict:
    with Image.open(path) as image:
        return {
            "path": str(path.relative_to(ROOT)),
            "width": image.width,
            "height": image.height,
            "format": image.format,
        }


def video_info(path: Path) -> dict:
    command = [
        "ffprobe",
        "-v",
        "error",
        "-select_streams",
        "v:0",
        "-show_entries",
        "stream=width,height,duration,avg_frame_rate",
        "-of",
        "json",
        str(path),
    ]
    try:
        result = subprocess.run(command, check=True, capture_output=True, text=True)
        return {"path": str(path.relative_to(ROOT)), **json.loads(result.stdout)}
    except (subprocess.CalledProcessError, FileNotFoundError, json.JSONDecodeError):
        return {"path": str(path.relative_to(ROOT)), "error": "ffprobe unavailable or failed"}


def markdown_outline(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    headings = re.findall(r"^(#{1,6})\s+(.+)$", text, flags=re.MULTILINE)
    images = re.findall(r"!\[[^\]]*\]\(([^)]+)\)|<img[^>]+src=[\"']([^\"']+)", text)
    code_blocks = re.findall(r"```([^\n]*)\n(.*?)```", text, flags=re.DOTALL)
    return {
        "path": str(path.relative_to(ROOT)),
        "characters": len(text),
        "headings": [{"level": len(mark), "text": title} for mark, title in headings],
        "images": [a or b for a, b in images],
        "code_blocks": [{"language": lang, "characters": len(code)} for lang, code in code_blocks],
    }


def main() -> None:
    result = {
        "template_docx": docx_text(ROOT / "《计算机应用项目实训-智能无人飞行器》设计报告模板.docx"),
        "acceptance_original": docx_text(ROOT / "验收过程记录表 .docx"),
        "acceptance_completed": docx_text(ROOT / "验收过程记录表（补充完整版）.docx"),
        "screenshots": [
            image_info(path)
            for path in sorted((ROOT / "实验截图").rglob("*"))
            if path.suffix.lower() in {".png", ".jpg", ".jpeg"}
        ],
        "videos": [
            video_info(path)
            for path in sorted((ROOT / "实验截图").rglob("*.mp4"))
        ],
        "courseware": [
            markdown_outline(path)
            for path in sorted((ROOT / "课件").glob("*.md"))
        ],
    }
    output = ROOT / "tmp" / "materials_inventory.json"
    output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(output)


if __name__ == "__main__":
    main()
