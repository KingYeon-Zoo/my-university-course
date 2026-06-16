from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/zoo/Desktop/计算机实训")
INPUT = ROOT / "tmp/智能无人飞行器设计报告-初版.docx"
OUTPUT = ROOT / "智能无人飞行器设计报告-朱清扬-2023212290.docx"
ARCH_IMAGE = ROOT / "tmp/系统总体架构图.png"
COVER_IMAGE = ROOT / "tmp/报告封面.png"
TABLE_IMAGE_DIR = ROOT / "tmp/报告表格"
LOGO = ROOT / "《计算机应用项目实训-智能无人飞行器》设计报告模板_images/image1.png"


def set_run_font(run, east_asia: str, latin: str, size: float, bold: bool | None = None) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "8")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])


def make_architecture_diagram() -> None:
    width, height = 1600, 820
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_paths = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    ]
    font_path = next(path for path in font_paths if Path(path).exists())
    title_font = ImageFont.truetype(font_path, 46)
    box_font = ImageFont.truetype(font_path, 31)
    small_font = ImageFont.truetype(font_path, 25)
    draw.text((width // 2, 35), "智能无人飞行器系统总体架构", fill="black", font=title_font, anchor="ma")

    boxes = {
        "camera": (70, 170, 340, 300, "机载摄像头"),
        "board": (430, 130, 780, 340, "泰山派 RK3566\nROS + YOLO/RKNN"),
        "flight": (430, 470, 780, 650, "飞控 + GPS\nMAVLink / MAVROS"),
        "target": (870, 140, 1210, 330, "目标检测与定位\n像素 + 姿态 + GPS"),
        "servo": (1290, 170, 1530, 300, "PWM 投弹舵机"),
        "vpn": (870, 500, 1210, 660, "4G + WireGuard\n10.0.0.0/24"),
        "ground": (1290, 470, 1530, 690, "阿里云服务器\n地面虚拟机\n实时图像与状态"),
    }

    for _, (x1, y1, x2, y2, label) in boxes.items():
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#EDF4FB", outline="#24557A", width=4)
        lines = label.split("\n")
        y = (y1 + y2) / 2 - (len(lines) - 1) * 22
        for line in lines:
            draw.text(((x1 + x2) / 2, y), line, fill="#102A43", font=box_font, anchor="mm")
            y += 47

    def arrow(start, end, label="") -> None:
        draw.line((start, end), fill="#333333", width=5)
        ex, ey = end
        sx, sy = start
        angle_x = 18 if ex >= sx else -18
        angle_y = 10 if ey >= sy else -10
        draw.polygon([(ex, ey), (ex - angle_x, ey - angle_y), (ex - angle_x, ey + angle_y)], fill="#333333")
        if label:
            draw.text(((sx + ex) / 2, (sy + ey) / 2 - 20), label, fill="#333333", font=small_font, anchor="mm")

    arrow((340, 235), (430, 235), "图像")
    arrow((780, 235), (870, 235), "检测结果")
    arrow((1210, 235), (1290, 235), "释放指令")
    arrow((605, 470), (605, 340), "飞行状态")
    arrow((780, 555), (870, 555), "任务数据")
    arrow((1040, 500), (1040, 330), "跨公网传输")
    arrow((1210, 580), (1290, 580), "VPN 中继")
    image.save(ARCH_IMAGE)


def make_cover_image() -> None:
    width, height = 1400, 1980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = next(
        path
        for path in (
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
        )
        if Path(path).exists()
    )

    def font(size: int) -> ImageFont.FreeTypeFont:
        return ImageFont.truetype(font_path, size)

    def centered(text: str, y: int, text_font, fill="black") -> None:
        box = draw.textbbox((0, 0), text, font=text_font)
        draw.text(((width - box[2] + box[0]) / 2, y), text, fill=fill, font=text_font)

    def centered_in_box(text: str, box, text_font, fill="black") -> None:
        x1, y1, x2, y2 = box
        bounds = draw.textbbox((0, 0), text, font=text_font)
        text_width = bounds[2] - bounds[0]
        text_height = bounds[3] - bounds[1]
        draw.text(
            ((x1 + x2 - text_width) / 2, (y1 + y2 - text_height) / 2 - bounds[1]),
            text,
            fill=fill,
            font=text_font,
        )

    logo = Image.open(LOGO).convert("RGBA")
    logo.thumbnail((760, 230), Image.Resampling.LANCZOS)
    image.paste(logo, ((width - logo.width) // 2, 105), logo)

    centered("《计算机应用项目实训》", 400, font(61))
    centered("智能无人飞行器设计报告", 490, font(61))
    centered("基于视觉识别、云端组网与自动投放的", 620, font(37), "#222222")
    centered("智能无人飞行器系统设计与实现", 680, font(37), "#222222")

    rows = [
        ("学生姓名", "朱清扬"),
        ("学号", "2023212290"),
        ("班级", "计科23-3班"),
        ("小组名称", "马刺总冠军"),
        ("个人分工", "目标识别、自动投放、云端组网与4G图传"),
        ("验收地点", "东操场"),
        ("完成日期", "2026年6月11日"),
    ]
    left, top = 175, 840
    table_width, row_height, label_width = 1050, 108, 300
    for index, (label, value) in enumerate(rows):
        y0 = top + index * row_height
        y1 = y0 + row_height
        draw.rectangle((left, y0, left + table_width, y1), outline="black", width=3)
        draw.line((left + label_width, y0, left + label_width, y1), fill="black", width=3)
        centered_in_box(label, (left, y0, left + label_width, y1), font(30))
        centered_in_box(value, (left + label_width, y0, left + table_width, y1), font(29))
    image.save(COVER_IMAGE, quality=96)


def wrap_table_text(draw, text: str, text_font, max_width: int) -> list[str]:
    lines: list[str] = []
    for source_line in (text or "").splitlines() or [""]:
        current = ""
        for char in source_line:
            candidate = current + char
            if current and draw.textlength(candidate, font=text_font) > max_width:
                lines.append(current)
                current = char
            else:
                current = candidate
        lines.append(current)
    return lines or [""]


def render_table_image(table, index: int) -> Path:
    TABLE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    column_count = len(table.columns)
    image_width = 2100 if column_count < 8 else 2500
    ratios_by_columns = {
        3: [0.20, 0.38, 0.42],
        4: [0.16, 0.22, 0.28, 0.34],
        9: [0.13, 0.09, 0.09, 0.075, 0.09, 0.075, 0.075, 0.12, 0.255],
    }
    ratios = ratios_by_columns.get(column_count, [1 / column_count] * column_count)
    column_widths = [int(image_width * ratio) for ratio in ratios]
    column_widths[-1] += image_width - sum(column_widths)
    font_path = next(
        path
        for path in (
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
        )
        if Path(path).exists()
    )
    body_font = ImageFont.truetype(font_path, 31 if column_count < 8 else 29)
    header_font = ImageFont.truetype(font_path, 32 if column_count < 8 else 29)
    measure = ImageDraw.Draw(Image.new("RGB", (1, 1)))
    padding_x, padding_y = 18, 15
    row_layouts = []
    for row_index, row in enumerate(table.rows):
        current_font = header_font if row_index == 0 else body_font
        cell_lines = []
        max_lines = 1
        for col_index, cell in enumerate(row.cells):
            lines = wrap_table_text(
                measure, cell.text.strip(), current_font, column_widths[col_index] - 2 * padding_x
            )
            cell_lines.append(lines)
            max_lines = max(max_lines, len(lines))
        line_height = current_font.size + 10
        row_height = max(72, max_lines * line_height + 2 * padding_y)
        row_layouts.append((cell_lines, row_height, current_font, line_height))

    image_height = sum(item[1] for item in row_layouts) + 4
    image = Image.new("RGB", (image_width + 4, image_height), "white")
    draw = ImageDraw.Draw(image)
    y = 2
    for row_index, (cell_lines, row_height, current_font, line_height) in enumerate(row_layouts):
        x = 2
        for col_index, lines in enumerate(cell_lines):
            width = column_widths[col_index]
            fill = "#D9EAF7" if row_index == 0 else "white"
            draw.rectangle((x, y, x + width, y + row_height), fill=fill, outline="black", width=3)
            text_height = len(lines) * line_height
            text_y = y + (row_height - text_height) / 2
            for line in lines:
                line_width = draw.textlength(line, font=current_font)
                draw.text(
                    (x + (width - line_width) / 2, text_y),
                    line,
                    fill="black",
                    font=current_font,
                )
                text_y += line_height
            x += width
        y += row_height
    path = TABLE_IMAGE_DIR / f"表格-{index}.png"
    image.save(path, dpi=(300, 300))
    return path


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag not in {qn("w:pPr")}:
            paragraph._p.remove(child)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, "宋体", "Times New Roman", 9)


def main() -> None:
    make_architecture_diagram()
    make_cover_image()
    doc = Document(INPUT)

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.3)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)
        footer = section.footer._element
        for child in list(footer):
            footer.remove(child)
        page_paragraph = section.footer.add_paragraph()
        add_page_number(page_paragraph)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_names = [("Heading 1", 16), ("Heading 2", 15), ("Heading 3", 13), ("Heading 4", 12)]
    for name, size in heading_names:
        if name not in styles:
            continue
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(8 if name == "Heading 1" else 5)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
    if "Heading 1" in styles:
        styles["Heading 1"].paragraph_format.page_break_before = False

    if "Source Code" in styles:
        code_style = styles["Source Code"]
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)
        code_style.paragraph_format.line_spacing = 1.0
        code_style.paragraph_format.first_line_indent = Pt(0)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)

    paragraphs = doc.paragraphs
    if len(paragraphs) < 5:
        raise RuntimeError("文档段落数量异常")

    cover_logo = paragraphs[0]
    clear_paragraph(cover_logo)
    cover_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_logo.paragraph_format.page_break_before = False
    cover_logo.paragraph_format.first_line_indent = Pt(0)
    cover_logo.paragraph_format.space_before = Pt(0)
    cover_logo.paragraph_format.space_after = Pt(0)
    cover_logo.add_run().add_picture(str(COVER_IMAGE), width=Inches(6.2))
    for paragraph in (paragraphs[1], paragraphs[2]):
        paragraph._element.getparent().remove(paragraph._element)

    for paragraph in paragraphs:
        text = paragraph.text.strip()
        paragraph.paragraph_format.widow_control = True

        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)

        if text.startswith("flowchart LR"):
            clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.add_run().add_picture(str(ARCH_IMAGE), width=Inches(6.2))

        if re.match(r"^(图|表)\s*[0-9A-Z]", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                set_run_font(run, "宋体", "Times New Roman", 10.5)

        if text.startswith("**关键词") or text.startswith("关键词："):
            paragraph.paragraph_format.first_line_indent = Pt(0)

    cover_table = doc.tables[0]
    cover_table._tbl.getparent().remove(cover_table._tbl)

    for table_index, table in enumerate(list(doc.tables), start=1):
        table_style = table._tbl.tblPr.find(qn("w:tblStyle"))
        if table_style is not None:
            table._tbl.tblPr.remove(table_style)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        column_count = len(table.columns)
        widths_by_columns = {
            3: [1900, 3300, 3300],
            4: [1300, 1900, 2400, 2900],
            9: [1100, 850, 850, 750, 850, 750, 750, 950, 1650],
        }
        widths = widths_by_columns.get(
            column_count, [8500 // column_count] * column_count
        )
        set_table_width(table, widths)

        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                tr_pr = row._tr.get_or_add_trPr()
                tbl_header = OxmlElement("w:tblHeader")
                tbl_header.set(qn("w:val"), "true")
                tr_pr.append(tbl_header)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if row_index == 0:
                    set_cell_shading(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            "宋体",
                            "Times New Roman",
                            7.5 if column_count >= 8 else 9.5,
                            bold=(row_index == 0),
                        )

    for table_index, table in enumerate(list(doc.tables), start=1):
        image_path = render_table_image(table, table_index)
        paragraph_element = OxmlElement("w:p")
        table._tbl.addprevious(paragraph_element)
        paragraph = Paragraph(paragraph_element, table._parent)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run().add_picture(str(image_path), width=Inches(6.2))
        table._tbl.getparent().remove(table._tbl)

    active_chapter = 1
    equation_counters: dict[int, int] = {}
    equation_counter = 0
    for paragraph in doc.paragraphs:
        heading_match = re.match(r"^(\d+)\.\s+", paragraph.text.strip())
        if paragraph.style and paragraph.style.name == "Heading 1" and heading_match:
            active_chapter = int(heading_match.group(1))
        if paragraph._p.xpath("./m:oMathPara"):
            equation_counters[active_chapter] = equation_counters.get(active_chapter, 0) + 1
            equation_counter += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.1), WD_TAB_ALIGNMENT.RIGHT)
            number_run = paragraph.add_run(
                f"\t({active_chapter}.{equation_counters[active_chapter]})"
            )
            set_run_font(number_run, "宋体", "Times New Roman", 10.5)

    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"tables={len(doc.tables)}, equations={equation_counter}, images={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()
